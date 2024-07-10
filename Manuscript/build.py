import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def add_style_from_source_to_destination(source_doc, destination_doc):
    """
    Copy missing styles from source to destination document.
    """
    source_styles = set(style.name for style in source_doc.styles)
    existing_styles = set(style.name for style in destination_doc.styles)
    
    for style_name in source_styles:
        if style_name not in existing_styles:
            # Create a new style in the destination with the same properties
            source_style = source_doc.styles[style_name]
            new_style = destination_doc.styles.add_style(style_name, source_style.type)
            new_style.base_style = source_style.base_style

def add_header_footer(doc, statement):
    """
    Adds a header and footer to the document.
    """
    header = doc.sections[0].header
    footer = doc.sections[0].footer
    
    header_text = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    footer_text = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    
    # Compile date
    compile_date = datetime.datetime.now().strftime('%Y-%m-%d')

    # Set header and footer content
    header_text.text = f"{statement} - {compile_date}"
    footer_text.text = f"{statement}"
    footer_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right-align the footer text

def merge_documents(file_list, output_file, statement, title, author):
    merged_document = Document()
    merged_document.core_properties.title = title
    merged_document.core_properties.author = author
    first_document = True

    for file in file_list:
        sub_doc = Document(file)
        if first_document:
            # Transfer all styles from the first document
            add_style_from_source_to_destination(sub_doc, merged_document)
            first_document = False
        else:
            # Optionally update styles for each subsequent document
            add_style_from_source_to_destination(sub_doc, merged_document)

        for paragraph in sub_doc.paragraphs:
            new_par = merged_document.add_paragraph(paragraph.text, style=paragraph.style.name)

        # Assume no page break needed after the last document
        if file != file_list[-1]:
            merged_document.add_page_break()

    # Add headers and footers
    add_header_footer(merged_document, statement)

    merged_document.save(output_file)

def main():
    with open('toc.yaml', 'r') as file:
        config = yaml.safe_load(file)

    file_list = config['files']
    output_file = config['output_file']
    statement = config['statement']
    title = config['title']
    author = config['author']
    merge_documents(file_list, output_file, statement, title, author)

if __name__ == "__main__":
    main()
